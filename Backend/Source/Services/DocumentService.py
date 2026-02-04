
import io
import docx
from fastapi import UploadFile, HTTPException
import fitz # PyMuPDF
import pandas as pd
import pytesseract
from PIL import Image
import pdf2image
import os
import shutil
from Backend.Source.Utils.FileValidator import FileValidator
from Backend.Source.Core.Logging import logger
from Backend.Source.Core.Config.Config import settings

class DocumentService:
    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        """
        Extracts text from various file formats (PDF, DOCX, XLSX, Images).
        Includes OCR fallback for scanned PDFs.
        Now with file validation.
        """
        # Read content
        content = await file.read()

        # VALIDATE FILE (NEW)
        try:
            safe_filename, file_size = await FileValidator.validate_upload(
                content,
                file.filename,
                max_size=settings.MAX_FILE_SIZE_CHAT  # Use chat-specific limit
            )
        except Exception as e:
            logger.error(f"File validation failed: {e}")
            raise

        # Use safe filename for processing
        filename = safe_filename.lower()

        logger.info(f"Processing file: {filename} ({file_size} bytes)")

        text = ""

        try:
            # --- PDF Handler ---
            if filename.endswith('.pdf'):
                with fitz.open(stream=content, filetype="pdf") as doc:
                    for page_num, page in enumerate(doc):
                        page_text = page.get_text()
                        if page_text.strip():
                            text += page_text + "\n"
                        else:
                            # Possible Image/Scan - Try OCR
                            logger.debug(f"Page {page_num} seems empty/scanned. Attempting OCR...")
                            try:
                                # Convert page to image (in memory)
                                pix = page.get_pixmap()
                                img_data = pix.tobytes("png")
                                img = Image.open(io.BytesIO(img_data))
                                # Run Tesseract
                                # Assumes Tesseract is in PATH or configured.
                                # If Arabic needed: lang='ara+eng'
                                ocr_text = pytesseract.image_to_string(img, lang='ara+eng')
                                text += f"\n[OCR Page {page_num}]:\n{ocr_text}\n"
                            except Exception as ocr_err:
                                logger.debug(f"OCR failed for page {page_num}: {ocr_err}")
                                text += "\n[Image/Scan - OCR Failed]\n"

            # --- Word Handler ---
            elif filename.endswith('.docx') or filename.endswith('.doc'):
                doc = docx.Document(io.BytesIO(content))
                # Extract Paragraphs
                para_text = "\n".join([para.text for para in doc.paragraphs])
                # Extract Tables (Important for Controls!)
                table_text = ""
                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_text += " | ".join(row_data) + "\n"
                
                text = f"{para_text}\n\n[TABLES]:\n{table_text}"

            # --- Excel Handler ---
            elif filename.endswith('.xlsx') or filename.endswith('.xls'):
                # Load workbook
                try:
                    df_dict = pd.read_excel(io.BytesIO(content), sheet_name=None) # Read all sheets
                    for sheet_name, df in df_dict.items():
                        text += f"\n\n--- Sheet: {sheet_name} ---\n"
                        # Convert to Markdown Table for LLM readability
                        text += df.to_markdown(index=False)
                except Exception as e:
                    logger.error(f"Excel read failed: {e}")
                    text = f"Error reading Excel file: {str(e)}"
            
            # --- Image Handler ---
            elif filename.endswith(('.png', '.jpg', '.jpeg')):
                try:
                    img = Image.open(io.BytesIO(content))
                    text = pytesseract.image_to_string(img, lang='ara+eng')
                except Exception as e:
                    text = f"Error performing OCR on image: {str(e)}"

            # --- Text Handler ---
            elif filename.endswith('.txt'):
                text = content.decode('utf-8')
            
            else:
                logger.warning(f"Unsupported file type: {filename}")
                return f"[Unsupported file type: {filename}]"

        except Exception as e:
            logger.error(f"Error parsing {filename}: {e}", exc_info=True)
            return f"[Error parsing file: {str(e)}]"
            
        return text

    @staticmethod
    def is_tesseract_installed():
        """Check if tesseract is available in PATH."""
        return shutil.which('tesseract') is not None
